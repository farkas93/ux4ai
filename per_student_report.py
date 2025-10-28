import json
import os
from collections import defaultdict
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import plotly.graph_objects as go
import plotly.io as pio

# Ensure Kaleido is available for image export
pio.kaleido.scope.default_format = "png"
pio.kaleido.scope.default_width = 700
pio.kaleido.scope.default_height = 400

# --- Language setup ---
LANG_DIR = "./lang"
DEFAULT_LANG = "en" # Default language for reports

def load_lang_file(lang_code):
    file_path = os.path.join(LANG_DIR, f"{lang_code}.json")
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            return json.load(f)
    except FileNotFoundError:
        print(f"Warning: Language file not found for '{lang_code}'. Falling back to default.")
        if lang_code == DEFAULT_LANG:
            raise FileNotFoundError(f"Default language file '{DEFAULT_LANG}.json' also not found.")
        return load_lang_file(DEFAULT_LANG)
    except json.JSONDecodeError:
        print(f"Error decoding JSON for '{lang_code}', falling back to default.")
        if lang_code == DEFAULT_LANG:
            raise json.JSONDecodeError(f"Default language file '{DEFAULT_LANG}.json' is malformed.", doc="", pos=0)
        return load_lang_file(DEFAULT_LANG)

# --- Directories ---
DATA_DIR = "./data"
SOLUTION_DIR = "./solutions"
REPORTS_DIR = "./reports" # New directory for generated reports

# Ensure reports directory exists
os.makedirs(REPORTS_DIR, exist_ok=True)

# --- Helper to load solution data ---
def load_solution_data(product_name):
    safe_filename = product_name.lower().replace(" ", "_").replace("/", "_") + ".json"
    solution_path = os.path.join(SOLUTION_DIR, safe_filename)
    if os.path.exists(solution_path):
        try:
            with open(solution_path, "r", encoding="utf-8") as f:
                return json.load(f)
        except (json.JSONDecodeError, IOError) as e:
            print(f"Error reading solution file {solution_path}: {e}")
    return None

# --- Function to create spider diagram for the report ---
def create_report_spider_diagram(user_scores, solution_scores, lang_dict, product_name):
    labels = ['conversational', 'specialization', 'autonomy', 'accessibility', 'explainability']
    # Extract just the main label part, removing descriptions in parentheses
    theta = [lang_dict[label + "_label"].split('(')[0].strip() for label in labels]
    theta_closed = theta + [theta[0]] # Close the loop

    fig = go.Figure()

    # User scores trace
    if user_scores:
        user_values = [user_scores.get(label.lower(), 0) for label in labels]
        user_values.append(user_values[0]) # Close the shape
        fig.add_trace(go.Scatterpolar(
            r=user_values,
            theta=theta_closed,
            fill='toself',
            name=lang_dict["user_answer_column_header"],
            line=dict(color='blue', width=3)
        ))

    # Solution scores trace
    if solution_scores:
        solution_values = [solution_scores.get(label.lower(), 0) for label in labels]
        solution_values.append(solution_values[0]) # Close the shape
        fig.add_trace(go.Scatterpolar(
            r=solution_values,
            theta=theta_closed,
            fill='toself',
            name=lang_dict["lecturer_answer_column_header"],
            line=dict(color='green', width=2)
        ))

    fig.update_layout(
        polar=dict(
            radialaxis=dict(visible=True, range=[0, 5])
        ),
        showlegend=True,
        title=f"{lang_dict['spider_diagram_label']} - {product_name}",
        height=400,
        margin=dict(l=50, r=50, t=70, b=50) # Adjust margins for better fit
    )
    return fig

# --- NEW: Function to create a plot for a single 0-5 value comparison ---
def create_single_value_comparison_plot(user_value, lecturer_value, title, min_label, max_label, lang_dict):
    fig = go.Figure()

    # Data for plotting
    y_labels = [lang_dict["user_answer_column_header"], lang_dict["lecturer_answer_column_header"]]
    x_values = [user_value if user_value is not None else 0, lecturer_value if lecturer_value is not None else 0]
    colors = ['blue', 'green']
    text_values = [f"{v:.1f}" if v is not None else "N/A" for v in [user_value, lecturer_value]]

    fig.add_trace(go.Bar(
        y=y_labels,
        x=x_values,
        orientation='h',
        marker_color=colors,
        text=text_values,
        textposition='outside',
        name='', # No legend needed for individual bars in this context
        width=0.8 # CHANGED: Make bars thinner
    ))

    fig.update_layout(
        title=title,
        xaxis=dict(
            range=[0, 5],
            tickvals=[0, 1, 2, 3, 4, 5],
            title_text=f"{min_label.split(':')[0].strip()} - {max_label.split(':')[0].strip()}",
            showgrid=True
        ),
        yaxis=dict(
            showgrid=False,
            categoryorder='array',
            categoryarray=y_labels # Ensure order is preserved
        ),
        height=100, # Compact height
        width=pio.kaleido.scope.default_width, # Ensure width is consistent with image export
        bargap=0.0, # CHANGED: Remove space between bars
        margin=dict(l=150, r=50, t=20, b=20), # Adjust margins
        showlegend=False
    )
    return fig


# --- Function to generate a report for a single student ---
def generate_student_report(username_full, student_product_files, lecturer_solutions, lang_dict):
    document = Document()
    document.add_heading(f"{lang_dict['report_title']} - {username_full}", level=1)

    for student_file_path in student_product_files:
        try:
            with open(student_file_path, "r", encoding="utf-8") as f:
                student_data = json.load(f)
        except (json.JSONDecodeError, IOError) as e:
            print(f"Error reading student file {student_file_path}: {e}")
            continue

        product_name = student_data.get("product_name", "Unknown Product")
        lecturer_data = lecturer_solutions.get(product_name, {}) # Ensure it's a dict even if not found

        # --- Product Name ---
        document.add_paragraph(f"{lang_dict['product_label']}: {product_name}", style='Heading 2')
        
        # --- AI Embedding ---
        document.add_paragraph(lang_dict['ai_embedding_label'])
        
        user_ai_role = student_data.get("ai_role", "N/A")
        lecturer_ai_role = lecturer_data.get("ai_role", "N/A")

        # CHANGED: Use a table for AI embedding
        table_ai_role = document.add_table(rows=1, cols=3)
        table_ai_role.style = 'Table Grid'
        hdr_cells_ai_role = table_ai_role.rows[0].cells
        hdr_cells_ai_role[0].text = ""
        hdr_cells_ai_role[1].text = lang_dict['user_answer_column_header']
        hdr_cells_ai_role[2].text = lang_dict['lecturer_answer_column_header']

        row_cells = table_ai_role.add_row().cells
        # Determine "Feature" or "Product" based on ai_role value
        if user_ai_role:
            user_display_role = lang_dict['ai_role_feature'] if "feature" in user_ai_role.lower() else lang_dict['ai_role_product']
        else:
            user_display_role = "N/A"
        lecturer_display_role = lang_dict['ai_role_feature'] if "feature" in lecturer_ai_role.lower() else lang_dict['ai_role_product']

        row_cells[0].text = lang_dict['ai_role_label'].split('(')[0].strip()
        row_cells[1].text = user_display_role
        row_cells[2].text = lecturer_display_role
        document.add_paragraph() # Spacer

        # --- UX4AI Spider Diagram ---
        document.add_heading(lang_dict['spider_diagram_label'], level=2)
        user_scores_for_plot = student_data.get('scores', {})
        solution_scores_for_plot = lecturer_data.get('scores', {})
        fig_spider = create_report_spider_diagram(user_scores_for_plot, solution_scores_for_plot, lang_dict, product_name)
        temp_spider_image_path = f"temp_spider_diagram_{username_full.replace(' ', '_')}_{product_name.replace(' ', '_')}.png"
        fig_spider.write_image(temp_spider_image_path)
        document.add_picture(temp_spider_image_path, width=Inches(6.5))
        os.remove(temp_spider_image_path)
        
        # --- Legend for Spider Diagram ---
        document.add_heading(lang_dict['legend_label'], level=3)
        legend_labels = [
            'conversational', 'specialization', 'autonomy', 'accessibility', 'explainability'
        ]
        for label_key in legend_labels:
            p = document.add_paragraph()
            p.add_run(lang_dict[label_key + "_label"])
        document.add_paragraph() # Spacer

        # --- Risk of Adversarial Attacks Section ---
        document.add_heading(lang_dict['risk_section_title_from_user'], level=2)
        document.add_paragraph(lang_dict['risk_level_info'])
        
        user_risk_level = student_data.get('risk_of_adversarial_attacks', {}).get('level')
        lecturer_risk_level = lecturer_data.get('risk_of_adversarial_attacks', {}).get('level')

        fig_risk = create_single_value_comparison_plot(
            user_risk_level, lecturer_risk_level,
            lang_dict['risk_level_label'].split('(')[0].strip(), # Title from label
            lang_dict['risk_level_label'].split('(')[1].split(',')[0].strip().replace(')', ''), # Min label
            lang_dict['risk_level_label'].split('(')[1].split(',')[1].strip().replace(')', ''), # Max label
            lang_dict
        )
        temp_risk_image_path = f"temp_risk_plot_{username_full.replace(' ', '_')}_{product_name.replace(' ', '_')}.png"
        fig_risk.write_image(temp_risk_image_path)
        document.add_picture(temp_risk_image_path, width=Inches(6.5))
        os.remove(temp_risk_image_path)
        document.add_paragraph() # Spacer

        # CHANGED: Table for only Risk Description
        table_risk_desc = document.add_table(rows=1, cols=3)
        table_risk_desc.style = 'Table Grid'
        hdr_cells_risk_desc = table_risk_desc.rows[0].cells
        hdr_cells_risk_desc[0].text = lang_dict['question_column_header']
        hdr_cells_risk_desc[1].text = lang_dict['user_answer_column_header']
        hdr_cells_risk_desc[2].text = lang_dict['lecturer_answer_column_header']

        user_risk_desc = student_data.get('risk_of_adversarial_attacks', {}).get('description', "N/A")
        lecturer_risk_desc = lecturer_data.get('risk_of_adversarial_attacks', {}).get('description', "N/A")
        
        row_cells = table_risk_desc.add_row().cells
        row_cells[0].text = lang_dict['risk_description_label'].split('(')[0].strip()
        row_cells[1].text = str(user_risk_desc)
        row_cells[2].text = str(lecturer_risk_desc)
        document.add_paragraph() # Spacer


        # --- Continuous Learning & Feedback Loops Section ---
        document.add_heading(lang_dict['continuous_learning_section_title_from_user'], level=2)
        document.add_paragraph(lang_dict['continuous_learning_aspects_info'])

        # CHANGED: Table for only Continuous Learning Aspects
        table_cl_aspects = document.add_table(rows=1, cols=3)
        table_cl_aspects.style = 'Table Grid'
        hdr_cells_cl_aspects = table_cl_aspects.rows[0].cells
        hdr_cells_cl_aspects[0].text = lang_dict['question_column_header']
        hdr_cells_cl_aspects[1].text = lang_dict['user_answer_column_header']
        hdr_cells_cl_aspects[2].text = lang_dict['lecturer_answer_column_header']

        cl_aspects_user = student_data.get('continuous_learning_feedback_loops', {}).get('aspects', "N/A")
        cl_aspects_lecturer = lecturer_data.get('continuous_learning_feedback_loops', {}).get('aspects', "N/A")
        
        row_cells = table_cl_aspects.add_row().cells
        row_cells[0].text = lang_dict['continuous_learning_aspects_label'].split('(')[0].strip()
        row_cells[1].text = str(cl_aspects_user)
        row_cells[2].text = str(cl_aspects_lecturer)
        document.add_paragraph() # Spacer

        # Analytics Type Level Plot
        document.add_paragraph(lang_dict['analytics_type_intro'])

        user_analytics_level = student_data.get('continuous_learning_feedback_loops', {}).get('analytics_type_level')
        lecturer_analytics_level = lecturer_data.get('continuous_learning_feedback_loops', {}).get('analytics_type_level')

        fig_analytics = create_single_value_comparison_plot(
            user_analytics_level, lecturer_analytics_level,
            lang_dict['analytics_type_label'].split('(')[0].strip(), # Title from label
            lang_dict['analytics_type_label'].split('(')[1].split(',')[0].strip().replace(')', ''), # Min label
            lang_dict['analytics_type_label'].split('(')[1].split(',')[1].strip().replace(')', ''), # Max label
            lang_dict
        )
        temp_analytics_image_path = f"temp_analytics_plot_{username_full.replace(' ', '_')}_{product_name.replace(' ', '_')}.png"
        fig_analytics.write_image(temp_analytics_image_path)
        document.add_picture(temp_analytics_image_path, width=Inches(6.5))
        os.remove(temp_analytics_image_path)
        document.add_paragraph() # Spacer

        # CHANGED: Table for only Analytics Explanation
        table_analytics_expl = document.add_table(rows=1, cols=3)
        table_analytics_expl.style = 'Table Grid'
        hdr_cells_analytics_expl = table_analytics_expl.rows[0].cells
        hdr_cells_analytics_expl[0].text = lang_dict['question_column_header']
        hdr_cells_analytics_expl[1].text = lang_dict['user_answer_column_header']
        hdr_cells_analytics_expl[2].text = lang_dict['lecturer_answer_column_header']

        analytics_explanation_user = student_data.get('continuous_learning_feedback_loops', {}).get('analytics_type_explanation', "N/A")
        analytics_explanation_lecturer = lecturer_data.get('continuous_learning_feedback_loops', {}).get('analytics_type_explanation', "N/A")
        
        row_cells = table_analytics_expl.add_row().cells
        row_cells[0].text = lang_dict['analytics_explanation_label'].split('(')[0].strip()
        row_cells[1].text = str(analytics_explanation_user)
        row_cells[2].text = str(analytics_explanation_lecturer)
        document.add_paragraph() # Spacer

        document.add_page_break()

    # Save the document
    safe_username = username_full.replace(" ", "_").replace("/", "_")
    report_filename = os.path.join(REPORTS_DIR, f"{safe_username}_UX4AI_Report.docx")
    document.save(report_filename)
    print(f"Report generated for {username_full}: {report_filename}")

# --- Main execution logic ---
if __name__ == "__main__":
    # You can change 'en' to 'de' or any other language code you have
    report_language = "de" # Changed to 'de' for demonstration with your provided German template
    current_lang_dict = load_lang_file(report_language)

    # Load all lecturer solutions once
    lecturer_solutions_by_product = {}
    for filename in os.listdir(SOLUTION_DIR):
        if filename.endswith(".json"):
            temp_product_name = filename.replace(".json", "").replace("_", " ").title()
            solution_data = load_solution_data(temp_product_name)
            if solution_data and solution_data.get("product_name"):
                 lecturer_solutions_by_product[solution_data["product_name"]] = solution_data
            elif solution_data:
                lecturer_solutions_by_product[temp_product_name] = solution_data

    # Group student data by username
    student_data_grouped = defaultdict(list)
    if os.path.exists(DATA_DIR):
        for user_dir in os.listdir(DATA_DIR):
            user_path = os.path.join(DATA_DIR, user_dir)
            if os.path.isdir(user_path):
                user_formatted_name_from_folder = user_dir
                full_username_for_report = user_dir

                for filename in os.listdir(user_path):
                    if filename.endswith(".json"):
                        file_path = os.path.join(user_path, filename)
                        student_data_grouped[user_formatted_name_from_folder].append(file_path)
                        
                        if full_username_for_report == user_dir:
                            try:
                                with open(file_path, "r", encoding="utf-8") as f:
                                    data = json.load(f)
                                    if data.get("username"): 
                                        full_username_for_report = data["username"] 
                            except:
                                pass

                if student_data_grouped[user_formatted_name_from_folder]:
                    student_data_grouped[full_username_for_report] = student_data_grouped.pop(user_formatted_name_from_folder)


    for username, product_files in student_data_grouped.items():
        if product_files:
            generate_student_report(username, product_files, lecturer_solutions_by_product, current_lang_dict)