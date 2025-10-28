import json
import os
from collections import defaultdict
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import plotly.graph_objects as go
import plotly.io as pio

# Ensure Kaleido is available for image export
# CHANGED: Updated to use plotly.io.defaults to address DeprecationWarning
pio.defaults.default_format = "png"
pio.defaults.default_width = 700
pio.defaults.default_height = 400

# --- Language setup ---
LANG_DIR = "./lang"
DEFAULT_LANG = "en" # Default language for reports

def load_lang_file(lang_code):
    file_path = os.path.join(LANG_DIR, f"{lang_code}.json")
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            return json.load(f)
    except FileNotFoundError:
        print(f"Warning: Language file not found for '{lang_code}'. Falling back to default.")
        if lang_code == DEFAULT_LANG:
            raise FileNotFoundError(f"Default language file '{DEFAULT_LANG}.json' also not found.")
        return load_lang_file(DEFAULT_LANG)
    except json.JSONDecodeError:
        print(f"Error decoding JSON for '{lang_code}', falling back to default.")
        if lang_code == DEFAULT_LANG:
            raise json.JSONDecodeError(f"Default language file '{DEFAULT_LANG}.json' is malformed.", doc="", pos=0)
        return load_lang_file(DEFAULT_LANG)

# --- Directories ---
DATA_DIR = "./data"
SOLUTION_DIR = "./solutions"
SUMMARY_REPORTS_DIR = "./summary_reports" # New directory for summary reports

# Ensure summary reports directory exists
os.makedirs(SUMMARY_REPORTS_DIR, exist_ok=True)

# --- Helper to load solution data ---
def load_solution_data(product_name):
    safe_filename = product_name.lower().replace(" ", "_").replace("/", "_") + ".json"
    solution_path = os.path.join(SOLUTION_DIR, safe_filename)
    if os.path.exists(solution_path):
        try:
            with open(solution_path, "r", encoding="utf-8") as f:
                return json.load(f)
        except (json.JSONDecodeError, IOError) as e:
            print(f"Error reading solution file {solution_path}: {e}")
    return None

# --- Function to create spider diagram for summary report ---
def create_summary_spider_diagram(avg_student_scores, lecturer_scores, lang_dict, product_name):
    labels = ['conversational', 'specialization', 'autonomy', 'accessibility', 'explainability']
    theta = [lang_dict[label + "_label"].split('(')[0].strip() for label in labels]
    theta_closed = theta + [theta[0]] # Close the loop

    fig = go.Figure()

    # Average Student scores trace
    if avg_student_scores:
        student_values = [avg_student_scores.get(label.lower(), 0) for label in labels]
        student_values.append(student_values[0]) # Close the shape
        fig.add_trace(go.Scatterpolar(
            r=student_values,
            theta=theta_closed,
            fill='toself',
            name=lang_dict["student_avg_label_plot"],
            line=dict(color='blue', width=3)
        ))

    # Lecturer scores trace
    if lecturer_scores:
        lecturer_values = [lecturer_scores.get(label.lower(), 0) for label in labels]
        lecturer_values.append(lecturer_values[0]) # Close the shape
        fig.add_trace(go.Scatterpolar(
            r=lecturer_values,
            theta=theta_closed,
            fill='toself',
            name=lang_dict["lecturer_ans_label_plot"],
            line=dict(color='green', width=2)
        ))

    fig.update_layout(
        polar=dict(
            radialaxis=dict(visible=True, range=[0, 5])
        ),
        showlegend=True,
        title=f"{lang_dict['spider_diagram_label']} - {product_name}",
        height=400,
        margin=dict(l=50, r=50, t=70, b=50)
    )
    return fig

# --- Function to create a plot for a single 0-5 value comparison (reused from individual report) ---
def create_single_value_comparison_plot(user_value, lecturer_value, title, min_label, max_label, lang_dict):
    fig = go.Figure()

    y_labels = [lang_dict["student_avg_label_plot"], lang_dict["lecturer_ans_label_plot"]]
    x_values = [user_value if user_value is not None else 0, lecturer_value if lecturer_value is not None else 0]
    colors = ['blue', 'green']
    text_values = [f"{v:.1f}" if v is not None else "N/A" for v in [user_value, lecturer_value]]

    fig.add_trace(go.Bar(
        y=y_labels,
        x=x_values,
        orientation='h',
        marker_color=colors,
        text=text_values,
        textposition='outside',
        name='',
        width=0.8
    ))

    fig.update_layout(
        title=title,
        xaxis=dict(
            range=[0, 5],
            tickvals=[0, 1, 2, 3, 4, 5],
            title_text=f"{min_label.split(':')[0].strip()} - {max_label.split(':')[0].strip()}",
            showgrid=True
        ),
        yaxis=dict(
            showgrid=False,
            categoryorder='array',
            categoryarray=y_labels
        ),
        height=200,
        width=pio.defaults.default_width, # CHANGED: Use pio.defaults.default_width
        bargap=0.0,
        margin=dict(l=150, r=50, t=50, b=50),
        showlegend=False
    )
    return fig


# --- Main aggregation and report generation logic ---
if __name__ == "__main__":
    report_language = "de" # Set desired language for the summary report
    current_lang_dict = load_lang_file(report_language)

    # --- 1. Load All Lecturer Solutions ---
    lecturer_solutions_by_product = {}
    for filename in os.listdir(SOLUTION_DIR):
        if filename.endswith(".json"):
            temp_product_name = filename.replace(".json", "").replace("_", " ").title()
            solution_data = load_solution_data(temp_product_name)
            if solution_data and solution_data.get("product_name"):
                 lecturer_solutions_by_product[solution_data["product_name"]] = solution_data
            elif solution_data:
                lecturer_solutions_by_product[temp_product_name] = solution_data

    # --- 2. Aggregate All Student Data ---
    aggregated_student_data = defaultdict(lambda: {
        "count": 0,
        "scores_sum": defaultdict(float),
        "ai_role_counts": defaultdict(int),
        "risk_level_sum": 0.0,
        "analytics_level_sum": 0.0,
        # For qualitative fields, we'll just note they are aggregated
        "risk_description_agg": [],
        "continuous_learning_aspects_agg": [],
        "analytics_explanation_agg": []
    })

    if os.path.exists(DATA_DIR):
        for user_dir in os.listdir(DATA_DIR):
            user_path = os.path.join(DATA_DIR, user_dir)
            if os.path.isdir(user_path):
                for filename in os.listdir(user_path):
                    if filename.endswith(".json"):
                        file_path = os.path.join(user_path, filename)
                        try:
                            # CHANGED: Corrected json.load() to open the file first
                            with open(file_path, "r", encoding="utf-8") as f:
                                student_data = json.load(f)
                            product_name = student_data.get("product_name")
                            if not product_name:
                                continue

                            agg = aggregated_student_data[product_name]
                            agg["count"] += 1

                            # Scores
                            if student_data.get("scores"):
                                for key, value in student_data["scores"].items():
                                    agg["scores_sum"][key] += value

                            # AI Role
                            if student_data.get("ai_role"):
                                agg["ai_role_counts"][student_data["ai_role"]] += 1

                            # Risk
                            if student_data.get("risk_of_adversarial_attacks"):
                                risk_info = student_data["risk_of_adversarial_attacks"]
                                agg["risk_level_sum"] += risk_info.get("level", 0.0)
                                if risk_info.get("description"):
                                    agg["risk_description_agg"].append(risk_info["description"])

                            # Continuous Learning
                            if student_data.get("continuous_learning_feedback_loops"):
                                cl_info = student_data["continuous_learning_feedback_loops"]
                                agg["analytics_level_sum"] += cl_info.get("analytics_type_level", 0.0)
                                if cl_info.get("aspects"):
                                    agg["continuous_learning_aspects_agg"].append(cl_info["aspects"])
                                if cl_info.get("analytics_explanation"):
                                    agg["analytics_explanation_agg"].append(cl_info["analytics_explanation"])

                        except (json.JSONDecodeError, IOError) as e:
                            print(f"Error reading student data file {file_path}: {e}")

    # Calculate averages from sums
    processed_student_data = {}
    for product, data in aggregated_student_data.items():
        count = data["count"]
        if count > 0:
            avg_scores = {k: v / count for k, v in data["scores_sum"].items()}
            processed_student_data[product] = {
                "count": count,
                "avg_scores": avg_scores,
                "ai_role_counts": data["ai_role_counts"],
                "avg_risk_level": data["risk_level_sum"] / count,
                "avg_analytics_level": data["analytics_level_sum"] / count,
                "risk_description_agg": data["risk_description_agg"],
                "continuous_learning_aspects_agg": data["continuous_learning_aspects_agg"],
                "analytics_explanation_agg": data["analytics_explanation_agg"]
            }

    # --- 3. Generate Summary DocX Report ---
    summary_document = Document()
    summary_document.add_heading(current_lang_dict['summary_report_title'], level=1)

    for product_name in sorted(processed_student_data.keys()):
        student_agg_data = processed_student_data.get(product_name)
        lecturer_data = lecturer_solutions_by_product.get(product_name, {})

        summary_document.add_heading(f"{current_lang_dict['product_label']}: {product_name}", level=2)

        if not student_agg_data:
            summary_document.add_paragraph(current_lang_dict['no_student_data_for_product'])
            summary_document.add_page_break()
            continue

        # Basic Summary
        summary_document.add_paragraph(f"{current_lang_dict['number_of_submissions_label']} {student_agg_data['count']}")
        
        # AI Role Distribution
        summary_document.add_paragraph(current_lang_dict['ai_role_distribution_label'])
        for role, count in student_agg_data['ai_role_counts'].items():
            summary_document.add_paragraph(f"- {role}: {count}")
        summary_document.add_paragraph() # Spacer

        # --- UX4AI Spider Diagram ---
        summary_document.add_heading(current_lang_dict['spider_diagram_label'], level=3)
        fig_spider = create_summary_spider_diagram(
            student_agg_data['avg_scores'],
            lecturer_data.get('scores', {}),
            current_lang_dict,
            product_name
        )
        temp_spider_image_path = f"temp_summary_spider_{product_name.replace(' ', '_')}.png"
        fig_spider.write_image(temp_spider_image_path)
        summary_document.add_picture(temp_spider_image_path, width=Inches(6.5))
        os.remove(temp_spider_image_path)
        summary_document.add_paragraph() # Spacer

        # --- AI Role Comparison Table ---
        summary_document.add_heading(current_lang_dict['ai_role_comparison_label'], level=3)
        table_ai_role = summary_document.add_table(rows=1, cols=3)
        table_ai_role.style = 'Table Grid'
        hdr_cells_ai_role = table_ai_role.rows[0].cells
        hdr_cells_ai_role[0].text = current_lang_dict['ai_role_label'].split('(')[0].strip()
        hdr_cells_ai_role[1].text = current_lang_dict['student_avg_label_plot']
        hdr_cells_ai_role[2].text = current_lang_dict['lecturer_ans_label_plot']

        row_cells = table_ai_role.add_row().cells
        # Determine most frequent student AI role
        most_frequent_student_role = "N/A"
        if student_agg_data['ai_role_counts']:
            most_frequent_student_role = max(student_agg_data['ai_role_counts'], key=student_agg_data['ai_role_counts'].get)
        
        lecturer_ai_role = lecturer_data.get('ai_role', "N/A")
        
        row_cells[0].text = ""
        row_cells[1].text = most_frequent_student_role
        row_cells[2].text = lecturer_ai_role
        summary_document.add_paragraph() # Spacer

        # --- Risk Level Comparison Plot ---
        summary_document.add_heading(current_lang_dict['risk_section_title_from_user'], level=3)
        fig_risk = create_single_value_comparison_plot(
            student_agg_data['avg_risk_level'],
            lecturer_data.get('risk_of_adversarial_attacks', {}).get('level'),
            current_lang_dict['avg_risk_level_label_short'],
            current_lang_dict['risk_level_label'].split('(')[1].split(',')[0].strip().replace(')', ''),
            current_lang_dict['risk_level_label'].split('(')[1].split(',')[1].strip().replace(')', ''),
            current_lang_dict
        )
        temp_risk_image_path = f"temp_summary_risk_plot_{product_name.replace(' ', '_')}.png"
        fig_risk.write_image(temp_risk_image_path)
        summary_document.add_picture(temp_risk_image_path, width=Inches(6.5))
        os.remove(temp_risk_image_path)
        summary_document.add_paragraph() # Spacer


        # --- Continuous Learning & Feedback Loops Section ---
        summary_document.add_heading(current_lang_dict['continuous_learning_section_title_from_user'], level=3)

        # --- Analytics Type Level Plot ---
        fig_analytics = create_single_value_comparison_plot(
            student_agg_data['avg_analytics_level'],
            lecturer_data.get('continuous_learning_feedback_loops', {}).get('analytics_type_level'),
            current_lang_dict['avg_analytics_level_label_short'],
            current_lang_dict['analytics_type_label'].split('(')[1].split(',')[0].strip().replace(')', ''),
            current_lang_dict['analytics_type_label'].split('(')[1].split(',')[1].strip().replace(')', ''),
            current_lang_dict
        )
        temp_analytics_image_path = f"temp_summary_analytics_plot_{product_name.replace(' ', '_')}.png"
        fig_analytics.write_image(temp_analytics_image_path)
        summary_document.add_picture(temp_analytics_image_path, width=Inches(6.5))
        os.remove(temp_analytics_image_path)
        summary_document.add_paragraph() # Spacer


        summary_document.add_page_break()

    # Save the summary document
    summary_report_filename = os.path.join(SUMMARY_REPORTS_DIR, f"UX4AI_Overall_Summary_Report_{report_language.upper()}.docx")
    summary_document.save(summary_report_filename)
    print(f"Overall summary report generated: {summary_report_filename}")